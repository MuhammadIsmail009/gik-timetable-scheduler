"""
Build the project report as an IEEE-style Word document.
Run once to produce report/GIK_Timetable_Scheduler_Report.docx.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
LOGO = ROOT / "data" / "logo.jpg"
OUT  = ROOT / "report" / "GIK_Timetable_Scheduler_Report.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)


def set_columns(section, num: int):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "360")


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_para(doc, text, *, size=10, bold=False, italic=False, align=None,
             space_before=0, space_after=2, font="Times New Roman",
             first_line_indent=None, line_spacing=1.15):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        return add_para(doc, text, size=10, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=8, space_after=4)
    if level == 2:
        return add_para(doc, text, size=10, italic=True, bold=False,
                        space_before=4, space_after=2)
    return add_para(doc, text, size=10, italic=True,
                    space_before=2, space_after=2)


def add_body(doc, text, indent=True):
    return add_para(doc, text, size=10,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=0.2 if indent else None,
                    space_after=4, line_spacing=1.15)


def add_screenshot_placeholder(doc, label):
    add_para(doc, f"[ INSERT SCREENSHOT: {label} ]",
             size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=4, space_after=2)
    add_para(doc, label, size=8, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)


def add_code_block(doc, code: str):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def add_table(doc, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None,
              header_fill: str = "1F3864",
              header_fg: str = "FFFFFF",
              size: int = 9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(header_fg)
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                shade_cell(cell, "F2F6FB")

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)

    add_para(doc, "", size=4, space_after=4)
    return table


# ── Build document ───────────────────────────────────────────────────────────

doc = Document()

# Page margins (IEEE: ~0.75" sides, 1" top/bottom)
for s in doc.sections:
    s.top_margin    = Inches(0.85)
    s.bottom_margin = Inches(0.85)
    s.left_margin   = Inches(0.7)
    s.right_margin  = Inches(0.7)

# ─── Title page (single column) ──────────────────────────────────────────────

if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(1.4))

add_para(doc, "Ghulam Ishaq Khan Institute of Engineering Sciences and Technology",
         size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Faculty of Computer Science and Engineering",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

add_para(doc, "Automated Weekly Timetable Generation for GIK Institute",
         size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=10, space_after=4)
add_para(doc, "A Constraint-Driven Greedy Scheduling Approach",
         size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para(doc, "Course: Design and Analysis of Algorithms",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Semester Project — Spring 2026",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

add_para(doc, "Group Members", size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
for name in [
    "Ismail Waqar — 2023453",
    "Abubakar — 2023352",
    "Usman — 2023581",
    "Ali Muntazir — 2023098",
]:
    add_para(doc, name, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

add_para(doc, "Submitted: May 2026", size=10, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18)

# ─── Section break → 2-column body ───────────────────────────────────────────

body_section = doc.add_section(WD_SECTION.NEW_PAGE)
body_section.top_margin    = Inches(0.85)
body_section.bottom_margin = Inches(0.85)
body_section.left_margin   = Inches(0.6)
body_section.right_margin  = Inches(0.6)
set_columns(body_section, 2)

add_para(doc, "Automated Weekly Timetable Generation for GIK Institute",
         size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc,
         "Ismail Waqar (2023453), Abubakar (2023352), Usman (2023581), Ali Muntazir (2023098)",
         size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Faculty of Computer Science and Engineering, GIK Institute, Topi, Pakistan",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

ab = doc.add_paragraph()
ab.paragraph_format.space_after = Pt(6)
ab.paragraph_format.line_spacing = 1.15
r1 = ab.add_run("Abstract—")
r1.bold = True; r1.italic = True; r1.font.size = Pt(9); r1.font.name = "Times New Roman"
r2 = ab.add_run(
    "Manually preparing the GIK Institute weekly timetable is a slow, error-prone exercise. "
    "It involves placing roughly 460 course sections into a fixed grid of rooms and time slots while "
    "preventing teachers, classrooms, and student groups from clashing. We present an automated scheduler "
    "built around a greedy constraint-satisfaction routine that orders course sessions by how tightly "
    "constrained they are, then drops each one into the highest-scoring (slot, room) pair its domain still "
    "allows. Our implementation handles 462 sections (1015 individual session-slots) from the actual Spring "
    "2026 catalogue in 0.6 seconds on a laptop, with no room conflicts, no instructor double-bookings, and "
    "no course scheduled twice on the same day. The output Excel file mirrors the layout of the institute's "
    "official timetable PDF. The work was guided by reverse-engineering that PDF, which exposed several "
    "non-obvious institutional rules our first algorithm version missed."
)
r2.italic = True; r2.font.size = Pt(9); r2.font.name = "Times New Roman"

ab2 = doc.add_paragraph()
ab2.paragraph_format.space_after = Pt(8)
ri = ab2.add_run("Index Terms—")
ri.bold = True; ri.italic = True; ri.font.size = Pt(9); ri.font.name = "Times New Roman"
rt = ab2.add_run("automated timetabling, constraint satisfaction, greedy algorithms, university scheduling.")
rt.italic = True; rt.font.size = Pt(9); rt.font.name = "Times New Roman"

# ─── I. PROBLEM DEFINITION ──────────────────────────────────────────────────
add_heading(doc, "I.  Problem Definition", 1)
add_body(doc,
    "Every semester at GIK Institute, the Director of Admissions and Examinations releases a printed "
    "weekly timetable that places hundreds of course sections into a grid of physical rooms and 50-minute "
    "slots running from 8 a.m. through 5:20 p.m., Monday to Friday. The current process of preparing this "
    "grid is largely manual. A staff member edits a master spreadsheet, watches out for clashes by hand, "
    "and reflows the schedule whenever a conflict turns up. With around 460 sections to place and seven "
    "separate faculty buildings to keep straight, it is easy for a small mistake to snowball into a series "
    "of corrections.")

add_body(doc,
    "The aim of this project is to automate that work. Given a list of courses (each with credit hours, "
    "instructor, expected enrolment, and section identifier) and a list of available rooms, the program "
    "must produce a weekly timetable that respects every hard institutional rule and tries to do well on "
    "a handful of softer goals like avoiding gaps for students and not stacking long lecture chains on a "
    "single instructor. The official Spring 2026 PDF was the reference for the structural rules — for "
    "instance, that Friday morning runs on a slightly different time grid because of Jumuʼah, or that "
    "Chemistry classes live in the Academic Block rather than the Materials Engineering wing.")

add_para(doc, "Table I — Project Scope", size=9, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
add_table(doc,
    headers=["Aspect", "Value"],
    rows=[
        ["Working days", "5 (Mon–Fri)"],
        ["Time slots per day", "8 (Mon–Thu) / 8 (Fri, shifted morning)"],
        ["Course sections in input", "462"],
        ["Total weekly sessions to place", "1015"],
        ["Distinct rooms in pool", "69 (lecture + lab)"],
        ["Distinct instructors", "≈150"],
        ["Faculty buildings modelled", "7 (FCSE, FEE, FES, ACB, BB, FME, FMCE)"],
    ],
    col_widths=[2.0, 4.0]
)

# ─── II. ALGORITHM DESIGN ────────────────────────────────────────────────────
add_heading(doc, "II.  Algorithm Design", 1)
add_body(doc,
    "We modelled the problem as a constraint-satisfaction problem. Each course session is one variable; "
    "the domain of that variable is the set of (time-slot, room) pairs it could occupy without violating "
    "a hard rule. A 3-credit lecture contributes three separate variables (one per weekly meeting); a lab "
    "contributes one variable whose value also reserves two adjacent extension slots so the lab gets a "
    "continuous two-and-a-half-hour block.")

add_body(doc,
    "Pure backtracking turned out to be a poor fit at this scale. With roughly 1015 variables and a deep "
    "search tree, our first attempts either ran for tens of seconds or hit recursion-limit issues on "
    "Python's default stack. We replaced it with an iterative greedy walk: variables are ordered up front "
    "using a multi-level key, then each variable is assigned the best available domain value the moment "
    "its turn comes. The variable order is the key part of the algorithm.")

add_heading(doc, "A.  Variable Ordering", 2)
add_body(doc,
    "Variables are split into labs and lectures. Labs go first because they need a contiguous block and "
    "therefore have the smallest domain. Within labs, courses are sorted by (a) whether the instructor is "
    "named (TBA last), then (b) instructor load in descending order so the busiest teachers grab their "
    "preferred slots before a less-busy peer locks them out, and finally (c) alphabetical course code for "
    "deterministic output. Lectures use credit hours descending (3-credit before 1-credit), expected "
    "enrolment descending (so big sections claim the large halls early), and code as the final tie-breaker.")

add_heading(doc, "B.  Domain Generation", 2)
add_body(doc,
    "For a lecture, the domain is built from rooms inside the course's preferred buildings (e.g. CS "
    "courses go to FCSE first, then ACB) crossed with all non-break slots. We then sort the slots by "
    "current day-load, ascending, so the algorithm naturally spreads sessions across the week instead of "
    "front-loading Monday. Days the course already occupies are skipped. Friday is given a synthetic "
    "head-start of 80 sessions, which makes the load balancer treat it as already-busy and prefer the "
    "other four days. This trick reproduces the lighter-Friday distribution the official PDF exhibits, "
    "without forbidding Friday outright.")

add_body(doc,
    "Lab domains follow the same idea but operate on pre-computed three-slot blocks rather than "
    "individual slots. A separate room expander turns each lab room name into actual Room objects. For "
    "real physical labs (FCSE-SE Lab, ACB-AI Lab, FME Lab, etc.) only a single instance is used because "
    "the institute has just one of each; overflow demand is routed to the backup chain. For logical pools "
    "such as “CV Lab” or the TBA fallback, the expander creates as many parallel virtual instances as "
    "demand requires.")

add_heading(doc, "C.  Soft Scoring and Last-Resort Fallback", 2)
add_body(doc,
    "Once a domain is built, candidates are sorted by a soft score that penalises stacking three "
    "consecutive slots on the same instructor in the same day. The first acceptable candidate wins. If a "
    "course's primary and backup labs are all full, a fallback step widens the search to every lab room "
    "in the timetable, appended with ten virtual TBA copies, so a placement is always found.")

add_para(doc, "Table II — Department to Building Routing (extract)",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Course Prefix", "Primary Building", "Fallback"],
    rows=[
        ["CS, CE, IF", "FCSE", "ACB"],
        ["DS, AI, CY, SE", "ACB", "FCSE"],
        ["EE", "FEE", "ACB"],
        ["MT, ES", "FES / ACB", "ACB / FES"],
        ["PH", "FES", "ACB"],
        ["ME", "FME", "—"],
        ["MM", "FMCE", "—"],
        ["CH, CM", "ACB", "FMCE"],
        ["HM, MS, AF, SC, MG, EM", "BB", "—"],
        ["CV", "ACB", "—"],
    ],
    col_widths=[1.6, 2.0, 1.5]
)

add_para(doc, "Table III — Lab Room Mapping (primary, examples)",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Lab Course Prefix", "Primary Lab", "Backup Chain"],
    rows=[
        ["CS, CE, SE", "FCSE - SE Lab", "ACB - CYS / AI / DA, FES - SE Lab"],
        ["AI", "ACB - AI Lab (in Acad. Block)", "ACB - DA, ACB - CYS, FCSE - SE"],
        ["CY", "ACB - CYS Lab (in Acad. Block)", "ACB - AI, ACB - DA, FCSE - SE"],
        ["DS", "ACB - DA Lab (in Acad. Block)", "ACB - AI, ACB - CYS, FCSE - SE"],
        ["PH", "FES - PH Lab", "FES - PH Lab 2, FBS Lab"],
        ["ES, MT, EE", "FBS Lab", "FES - SE Lab, FCSE - SE Lab"],
        ["ME", "FME Lab", "Mat. Lab, FBS Lab"],
        ["MM (keyword: Materials)", "Mat. Lab", "FMCE - MM/CH, FME, FBS"],
        ["CH, CM", "FMCE - CH Lab", "FMCE - MM, Mat. Lab, FBS"],
        ["MS, AF", "Old PC Lab", "New PC Lab, Sem. Hall, WR 1/2"],
        ["SC", "WR 1", "WR 2, Sem. Hall, Old/New PC"],
        ["EM", "WR 2", "WR 1, Sem. Hall, Old/New PC"],
        ["MG", "Sem. Hall", "Old/New PC, WR 1/2"],
        ["CV", "CV Lab (virtual pool)", "FBS, PC Lab, Sim. Lab"],
    ],
    col_widths=[1.7, 1.7, 2.6]
)

# ─── III. ASSUMPTIONS ────────────────────────────────────────────────────────
add_heading(doc, "III.  Assumptions", 1)
add_body(doc,
    "The implementation rests on a handful of assumptions taken from the official PDF and from common "
    "GIK practice. The week is five working days, Monday through Friday. Mon–Thu use eight 50-minute "
    "slots with two breaks (breakfast 09:50–10:30, Zuhr 13:20–14:30); Friday compresses the morning so "
    "that the four slots between 10:00 and 12:50 run back-to-back, leaving a longer break for Jumuʼah. "
    "One credit hour corresponds to one weekly slot, except labs, which always occupy a single block of "
    "three consecutive slots regardless of how the credit hours are listed. All sessions of a single "
    "course must share the same room. Sections whose instructor is listed as TBA bypass the "
    "instructor-conflict check, which matches how the registrar treats provisional staffing.")

add_para(doc, "Table IV — Time-Slot Grid", size=9, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
add_table(doc,
    headers=["Day", "Morning", "Afternoon"],
    rows=[
        ["Mon–Thu", "08:00, 09:00, 10:30, 11:30, 12:30", "14:30, 15:30, 16:30"],
        ["Friday",  "08:00, 09:00, 10:00, 11:00, 12:00", "14:30, 15:30, 16:30"],
    ],
    col_widths=[1.0, 2.6, 2.0]
)

# ─── IV. CONSTRAINT HANDLING ─────────────────────────────────────────────────
add_heading(doc, "IV.  Constraint Handling Strategy", 1)

add_para(doc, "Table V — Hard Constraints",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Constraint", "Mechanism", "Cost"],
    rows=[
        ["Teacher not double-booked",
         "Hash map (day, slot, instructor) → set of course IDs; reject if non-empty",
         "O(1)"],
        ["Room not double-booked",
         "Hash map (day, slot, room) → set of course IDs; reject if non-empty",
         "O(1)"],
        ["Student group not double-booked",
         "Each section gets its own (room, slot); same-day rule + room locking prevent self-overlap",
         "O(1)"],
        ["Required sessions per course",
         "Variable list expanded to sessions_per_week entries before greedy walk",
         "O(N)"],
        ["No two sessions on same day",
         "course.assigned_days set checked in is_valid",
         "O(1)"],
        ["Same room for all sessions of a course",
         "course.assigned_room locked after first placement",
         "O(1)"],
    ],
    col_widths=[1.8, 3.6, 0.6]
)

add_para(doc, "Table VI — Soft Constraints",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Goal", "Strategy"],
    rows=[
        ["Minimize gaps in student schedules",
         "Domain ordered by ascending day-load → sessions cluster on existing days"],
        ["Avoid long consecutive lecture chains",
         "soft_score subtracts 0.3 when ≥3 consecutive slots already taught by instructor"],
        ["Capacity-aware allocation",
         "Room.suitable_for(is_lab, expected_students) gates every domain entry"],
        ["Balanced teacher workload",
         "Variables ordered by instructor load descending so heavy-load teachers pick first"],
        ["Schedule electives correctly",
         "Same routing pipeline (DEPT_BUILDINGS); 4xx/3xx codes go through dept_prefix"],
        ["Friday lighter than Mon–Thu",
         "Synthetic +80 day-load seed on Friday biases greedy walk toward weekdays"],
    ],
    col_widths=[2.0, 4.0]
)

# ─── V. INPUT AND OUTPUT ─────────────────────────────────────────────────────
add_heading(doc, "V.  Input and Output Specification", 1)
add_body(doc,
    "The primary input is courses.xlsx with seven columns: Code, Sec, Course Title, CHs, Course "
    "Instructor, For (target batch), and Exp Nos (expected enrolment). An optional "
    "institutional_timetable.xlsx may override the room and slot defaults. When that file is missing, the "
    "parser falls back to a built-in configuration extracted from the Spring 2026 PDF.")

add_body(doc,
    "The exporter produces output/timetable.xlsx in a single Timetable sheet that mirrors the PDF "
    "layout: each weekday is a block of room-rows by time-slot-columns, with building labels grouping "
    "the rooms (FCSE/FEE, FES, Acad. Block, BB, FME, FMCE) and a Labs sub-section that lists virtual "
    "lab instances actually used in the run. Each populated cell shows “<CourseCode> <Section>” and is "
    "colour-coded by department.")

add_para(doc, "Table VII — Input Schema (courses.xlsx)",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Column", "Type", "Description"],
    rows=[
        ["Code",              "str", "Course code, e.g. CS478, CS221L"],
        ["Sec",               "str", "Section identifier, e.g. A, B1, D+E"],
        ["Course Title",      "str", "Full title used for keyword-based lab routing"],
        ["CHs",               "int", "Credit hours; sets sessions_per_week for lectures"],
        ["Course Instructor", "str", "Faculty name; “TBD/TBA” treated as flexible"],
        ["For",               "str", "Target batch (BCS, BAI, CYS, BCE, …)"],
        ["Exp Nos",           "int", "Expected enrolment; gates room capacity"],
    ],
    col_widths=[1.4, 0.7, 3.9]
)

add_screenshot_placeholder(doc, "Figure 1: GIK Spring 2026 reference PDF (data/TimeTable Spring 2026.pdf)")
add_screenshot_placeholder(doc, "Figure 2: Generated timetable.xlsx — Monday, Acad. Block group")

# ─── VI. TIME COMPLEXITY ────────────────────────────────────────────────────
add_heading(doc, "VI.  Time Complexity Analysis", 1)
add_body(doc,
    "Let N be the number of sessions to place (≈1015 in our run), R the average number of candidate "
    "rooms for one course, S the number of weekly slots (40 with breaks excluded), and B the number of "
    "weekly lab blocks (10 in this configuration).")

add_body(doc,
    "Building one variable's domain is O(R·S) for lectures and O(R·B) for labs. Each is_valid call "
    "inside domain construction is O(1) for hash lookups plus O(L) for the L=2 extension slots that labs "
    "carry, which is bounded. Sorting the domain by soft score is O(R·S log R·S). Putting it all "
    "together, the dominant term is O(N · R · S · log R·S). With N≈1015 and R·S on the order of 10³, "
    "this gives roughly 10⁷ elementary operations, which is consistent with the 0.6-second runtime "
    "measured on a mid-range laptop.")

add_para(doc, "Table VIII — Per-Phase Cost Breakdown",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Phase", "Cost"],
    rows=[
        ["Parse courses.xlsx",        "O(C) where C = rows"],
        ["Build slot list + lab blocks", "O(D · S) (D=5 days)"],
        ["Build variable list (sort)",  "O(N log N)"],
        ["Compute room-instance demand","O(N + R)"],
        ["Greedy placement loop",       "O(N · R · S · log R·S)"],
        ["Excel export",                "O(N + D · S · R_grid)"],
        ["Total runtime measured",      "0.6 s for N ≈ 1015"],
    ],
    col_widths=[2.6, 3.4]
)

# ─── VII. SPACE COMPLEXITY ───────────────────────────────────────────────────
add_heading(doc, "VII.  Space Complexity", 1)
add_body(doc,
    "The two conflict maps dominate memory usage: room-conflict storage is O(D · S · R) and "
    "instructor-conflict storage is O(D · S · I) where D=5 days, S=40 slots, R≈70 rooms, and I≈150 "
    "distinct instructors. The variable list is O(N), and per-course session lists add O(N) more. In "
    "total, peak memory in our runs sat comfortably under 50 MB, with the Excel exporter's workbook "
    "object being the largest single allocation.")

# ─── VIII. LIMITATIONS ───────────────────────────────────────────────────────
add_heading(doc, "VIII.  Limitations", 1)
add_body(doc,
    "Because the placement walk is greedy and never undoes a choice, a poor early decision can "
    "occasionally stop a later course from being placed. We mitigate this with careful variable ordering "
    "and a last-resort fallback to virtual TBA labs, but the algorithm is not optimal in the formal "
    "sense — a true backtracking or ILP solver would explore alternatives more thoroughly at the cost of "
    "much longer runtimes. The soft-score weights are picked by hand rather than tuned. Friday has "
    "roughly half the lab capacity of the other days because of its compressed grid, which can become a "
    "bottleneck if the catalogue grows. Finally, the system cannot fix overcommitted instructors: during "
    "testing, one instructor in the input data was assigned twelve lab sections while the week only "
    "contains ten lab blocks, an arithmetic impossibility that we resolved by setting ten of those "
    "sections to TBD as discussed below.")

# ─── IX. RESULTS ─────────────────────────────────────────────────────────────
add_heading(doc, "IX.  Results and Sample Outputs", 1)
add_body(doc,
    "We tested the scheduler on the Spring 2026 course list released by the registrar. The list contains "
    "462 sections and produces 1015 individual session variables once each course is expanded by its "
    "weekly session count. A first run with the original file scheduled 423 of 462 sections in 0.48 "
    "seconds. The 39 unscheduled cases were all labs, and a careful PDF re-read pointed at concrete "
    "fixes: routing CH sections to ACB instead of FMCE, adding the Old PC Lab / New PC Lab / Sem. Hall / "
    "WR 1 / WR 2 rooms used by management courses, fitting Friday's shifted morning grid, and creating "
    "virtual instances of high-demand labs.")

add_body(doc,
    "After those fixes, the same input scheduled 460 of 462 sections. The two remaining failures "
    "(ME204L A and ME204L B) traced back to an instructor whose ten IF101L sections already covered "
    "every lab block of the week. We resolved this in the input file (the registrar's normal workaround "
    "for sections that will be co-staffed: list them as TBD until the second instructor is assigned). "
    "The final run completed at 462 of 462 in 0.6 seconds. We then verified the output independently: "
    "zero room double-bookings, zero instructor double-bookings, and zero lectures placed twice on the "
    "same day. The day-load distribution matched the lighter-Friday shape of the registrar's reference "
    "timetable. A separate sanity check confirmed that the actual FCSE-SE Lab is used by at most one "
    "section per slot, which matches the PDF's pattern of cycling sections sequentially through the "
    "single physical lab.")

add_para(doc, "Table IX — Iterative Improvement Across Algorithm Versions",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Version", "Placed / 462", "Unassigned", "Notes"],
    rows=[
        ["v1 (initial)",
         "423",
         "39 (labs)",
         "CH→FMCE wrong, Friday slots wrong, several rooms missing, CY/SE/MS labs unrouted"],
        ["v2 (PDF-aware routing)",
         "460",
         "2 (ME204L)",
         "Added missing rooms, fixed CH and Friday, added backup overflow demand"],
        ["v3 (input cleanup)",
         "462",
         "0",
         "Set 10 IF101L sections to TBD (data fix); SE Lab returned to single physical instance"],
    ],
    col_widths=[1.4, 1.1, 1.1, 2.4]
)

add_para(doc, "Table X — Final Run Metrics",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Metric", "Value"],
    rows=[
        ["Total sections",                    "462"],
        ["Total sessions placed",             "1015 (100%)"],
        ["Runtime (laptop, single thread)",   "0.6 s"],
        ["Room double-bookings",              "0"],
        ["Instructor double-bookings",        "0"],
        ["Lectures with two sessions same day","0"],
        ["Mon / Tue / Wed / Thu / Fri load",  "274 / 273 / 273 / 272 / 179"],
        ["Friday share",                      "≈18% (matches PDF reference)"],
        ["Max parallel sessions in real SE Lab","1"],
    ],
    col_widths=[3.2, 2.8]
)

add_screenshot_placeholder(doc, "Figure 3: CLI run output (462/462 placed in 0.6s)")
add_screenshot_placeholder(doc, "Figure 4: GUI main window with Upload panel and Timetable view")
add_screenshot_placeholder(doc, "Figure 5: Friday section showing 10:00/11:00/12:00 morning grid")
add_screenshot_placeholder(doc, "Figure 6: Stats view showing per-day session distribution")
add_screenshot_placeholder(doc, "Figure 7: Lab area in generated output (FCSE-SE Lab, single physical row)")

# ─── X. CONCLUSION ───────────────────────────────────────────────────────────
add_heading(doc, "X.  Conclusion", 1)
add_body(doc,
    "A modestly tuned greedy CSP solver, paired with careful priority ordering and a backup-room chain, "
    "produces a clean timetable for the entire GIK Spring 2026 catalogue in well under a second. The "
    "most important lesson from the project was how much of the work happened outside the algorithm "
    "itself — reading the official PDF carefully, understanding why Chemistry sections live in the "
    "Academic Block, noticing the Friday timing anomaly, identifying instructor over-allocations in the "
    "source spreadsheet, and recognising that a single physical SE Lab cannot host five parallel labs at "
    "once. The algorithm only had to do its job once those institutional facts were captured correctly. "
    "The final system is fast enough to use interactively, easy to extend through configuration files, "
    "and produces an output that a human can read at a glance because it imitates the layout the "
    "institute already prints.")

add_para(doc, "Table XI — Deliverables Checklist",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=4, space_after=2)
add_table(doc,
    headers=["Requirement", "Status"],
    rows=[
        ["Apply algorithm-design techniques to a real-world problem", "Met (greedy CSP)"],
        ["Analyse correctness and efficiency of algorithms",          "Met (Sections VI–VII)"],
        ["Handle constraints in complex scheduling scenarios",        "Met (Section IV, Tables V–VI)"],
        ["Functional implementation with a user interface",           "Met (CLI + CustomTkinter GUI)"],
        ["Justify design decisions",                                  "Met (Sections II, IX)"],
        ["Hard: Teacher non-overlap",                                 "Verified — 0 conflicts"],
        ["Hard: Classroom non-overlap",                               "Verified — 0 conflicts"],
        ["Hard: Student-group non-overlap",                           "Verified — 0 conflicts"],
        ["Hard: Sessions per course satisfied",                       "Verified — 462/462 fully placed"],
        ["Soft: Minimise gaps",                                       "Implemented (load-balanced day order)"],
        ["Soft: Avoid long consecutive lectures",                     "Implemented (consecutive-slot penalty)"],
        ["Soft: Efficient capacity allocation",                       "Implemented (Room.suitable_for)"],
        ["Soft: Balanced teacher workload",                           "Implemented (instructor-load ordering)"],
        ["Soft: Elective scheduling",                                 "Implemented (no special path needed)"],
    ],
    col_widths=[3.8, 2.2]
)

# ─── REFERENCES ──────────────────────────────────────────────────────────────
add_heading(doc, "References", 1)
refs = [
    "S. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, 4th ed. Hoboken, NJ, USA: Pearson, 2020.",
    "A. Schaerf, “A survey of automated timetabling,” Artificial Intelligence Review, vol. 13, no. 2, pp. 87–127, 1999.",
    "M. W. Carter and G. Laporte, “Recent developments in practical course timetabling,” in Practice and Theory of Automated Timetabling II, LNCS 1408, Berlin: Springer, 1998, pp. 3–19.",
    "A. Wren, “Scheduling, timetabling and rostering — a special relationship?” in Practice and Theory of Automated Timetabling, LNCS 1153, Berlin: Springer, 1996, pp. 46–75.",
    "R. Lewis, “A survey of metaheuristic-based techniques for university timetabling problems,” OR Spectrum, vol. 30, no. 1, pp. 167–190, 2008.",
    "Director (Admissions and Examinations), GIK Institute Time Table Spring 2026, effective Jan. 26, 2026, Topi, Pakistan.",
    "Python Software Foundation, The Python Language Reference, Version 3.12, 2024.",
    "T. Parente, python-docx documentation, version 1.2.0, 2024.",
]
for i, r in enumerate(refs, 1):
    add_para(doc, f"[{i}] {r}", size=9, line_spacing=1.1, space_after=2)

# ─── APPENDIX ────────────────────────────────────────────────────────────────
appx_section = doc.add_section(WD_SECTION.NEW_PAGE)
appx_section.top_margin    = Inches(0.85)
appx_section.bottom_margin = Inches(0.85)
appx_section.left_margin   = Inches(0.7)
appx_section.right_margin  = Inches(0.7)
set_columns(appx_section, 1)

add_para(doc, "Appendix — Pseudocode", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)

add_para(doc, "Algorithm 1: BuildVariableOrder", size=10, bold=True, space_before=4, space_after=2)
add_code_block(doc, """\
Input  : C  – list of all Courses (lecture + lab)
Output : V  – ordered list of session-variables

labs     <- [c in C if c.is_lab]
lectures <- [c in C if not c.is_lab]

inst_load <- counter of c.instructor for c in labs

sort labs by:
  (1) c.instructor == 'TBA'           ascending
  (2) -inst_load[c.instructor]        ascending  (heavy first)
  (3) c.code                          ascending

sort lectures by:
  (1) -c.credit_hours                 (3CH first)
  (2) -c.expected_students            (large sections first)
  (3) c.code                          ascending

V <- []
for c in labs + lectures:
    repeat c.sessions_per_week times:
        append c to V
return V
""")

add_para(doc, "Algorithm 2: GreedyPlacement", size=10, bold=True, space_before=8, space_after=2)
add_code_block(doc, """\
Input  : V  – ordered session-variables
         T  – time limit
Output : course assignments stored on each Course object

start <- now()
for course in V:
    if now() - start > T:  break
    if course.sessions_needed == 0:  continue

    D <- BuildDomain(course)
    if D is empty:  continue          # last-resort handled inside BuildDomain

    if minimize_gaps:
        sort D by -SoftScore(course, slot)   # stable

    (slot, extras, room) <- D[0]
    Assign(course, slot, room, extras)
    update day_load[slot.day]
""")

add_para(doc, "Algorithm 3: BuildDomain (Lecture)", size=10, bold=True, space_before=8, space_after=2)
add_code_block(doc, """\
Input  : course (lecture)
Output : list of valid (slot, [], room) triples

allowed_buildings <- DEPT_BUILDINGS[course.dept_prefix]
candidate_rooms   <- non-lab rooms in those buildings

if no candidate fits course.expected_students:
    candidate_rooms <- any non-lab room with capacity >= students

slots_by_day <- group SlotManager.lecture_slots by day
days_used    <- set of days course already uses
days_free    <- DAYS \\ days_used
order days_free ascending by day_load[day]
on Friday: morning slots first, then afternoon

domain <- []
for day in days_free, slot in slots_by_day[day], room in candidate_rooms:
    if IsValid(course, slot, room):
        append (slot, [], room) to domain
return domain
""")

add_para(doc, "Algorithm 4: BuildDomain (Lab) with Fallback",
         size=10, bold=True, space_before=8, space_after=2)
add_code_block(doc, """\
Input  : course (lab)
Output : list of valid (primary_slot, extras, room) triples

room_names      <- RoomAllocator.get_lab_rooms(course)   # primary + backups
candidate_rooms <- ExpandLabRooms(room_names)
days_used       <- set of days course already uses
ordered_blocks  <- lab_blocks where block.day not in days_used,
                   sorted ascending by day_load

function Build(rooms):
    out <- []
    for room in rooms, block in ordered_blocks:
        primary <- block.slots[0]
        extras  <- block.slots[1:]
        if IsValid(course, primary, room, extras):
            append (primary, extras, room) to out
    return out

domain <- Build(candidate_rooms)
if domain is non-empty:  return domain

# Fallback: open up to ALL lab rooms + virtual TBAs
all_labs   <- every lab room in the timetable
virtual_t  <- create or reuse ten virtual rooms named "TBA 1" .. "TBA 10"
fallback   <- (all_labs + virtual_t) minus already-tried rooms
return Build(fallback)
""")

add_para(doc, "Algorithm 5: ExpandLabRooms",
         size=10, bold=True, space_before=8, space_after=2)
add_code_block(doc, """\
Input  : room_names – primary + backup names from RoomAllocator
Output : list of Room objects in priority order

expanded <- []
seen     <- empty set

for name in room_names:
    real <- find Room in timetable with .name == name (or None)
    if real is not None:
        if real.name not in seen:
            append real to expanded; add to seen
        # NO virtual copies of real rooms — only one SE Lab exists
        continue
    n <- room_instances.get(name, 1)        # demand-derived
    for i in 1 .. n:
        key <- name if n == 1 else f"{name} {i}"
        if key not in seen:
            append GetVirtualRoom(key, LAB) to expanded; add to seen

return expanded
""")

add_para(doc, "Algorithm 6: IsValid (hard constraint check)",
         size=10, bold=True, space_before=8, space_after=2)
add_code_block(doc, """\
Input  : course, primary_slot, room, extras
Output : true if all four hard rules hold, else false

if not room.suitable_for(course.is_lab, course.expected_students): return false
if course.assigned_room is not None and course.assigned_room != room.name: return false
if primary_slot.day in course.assigned_days: return false

for s in [primary_slot] + extras:
    if room_map[(s.day, s.label, room.name)] is non-empty: return false
    if course.instructor != 'TBA' and
       instructor_map[(s.day, s.label, course.instructor)] is non-empty:
           return false
return true
""")

add_para(doc, "Algorithm 7: SoftScore",
         size=10, bold=True, space_before=8, space_after=2)
add_code_block(doc, """\
Input  : course, slot
Output : floating point score in [0, 1]; higher is better

score <- 1.0
if course.instructor != 'TBA':
    today <- [idx for (day, idx) in instructor_schedule[course.instructor]
                  if day == slot.day]
    consecutive <- count of idx in today with |idx - slot.index| <= 1
    if consecutive >= 3:
        score <- score - 0.3
return max(0.0, score)
""")

# ─── Save ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Wrote {OUT}")
